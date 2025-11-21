"""Генерация подробного руководства в формате .docx (рус.)
Требует установленный пакет python-docx.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    return h


def add_par(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def add_list(doc: Document, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def main():
    doc = Document()

    # Базовая настройка шрифта/языка
    styles = doc.styles
    style = styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    title = doc.add_heading('ЭкоБарахолка — Руководство по установке и использованию', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_par(doc, f"Версия документа: {datetime.now():%Y-%m-%d %H:%M}")

    add_heading(doc, '1. Назначение и краткое описание', 1)
    add_par(doc, 'Проект "ЭкоБарахолка" — учебное приложение для обмена вещами, пожертвований и ответственной утилизации. Включает веб-интерфейс (Flask), REST API и в дальнейшем десктоп-клиент PySide6.',)

    add_heading(doc, '2. Архитектура и структура проекта', 1)
    add_par(doc, 'Проект разделён на слои, чтобы код был поддерживаемым и расширяемым:')
    add_list(doc, [
        'backend/app — Flask-приложение: модели, HTML-маршруты, REST API, шаблоны, статика',
        'backend/repositories — слой доступа к данным (DAL), будет наполняться',
        'backend/services — бизнес-логика, будет наполняться',
        'client — задел для десктоп-клиента на PySide6',
        'docs — документация (архитектура, гайд, тест-план, тест-кейсы)',
        'scripts — служебные скрипты (инициализация БД, сиды, архив, генерация .docx)',
        'data — дампы и тестовые данные (по мере наполнения)'
    ])

    add_heading(doc, '3. Требования к окружению', 1)
    add_list(doc, [
        'Windows 10/11, PowerShell',
        'Python 3.8+ (рекомендуется 3.10+)',
        'MySQL 8.x (или PostgreSQL при необходимости)',
        'pip для установки зависимостей',
    ])

    add_heading(doc, '4. Установка и запуск (MySQL)', 1)
    add_par(doc, 'Шаг 1. Создайте базу и пользователя в MySQL (через MySQL Workbench либо SQL):', True)
    add_par(doc, 'CREATE DATABASE ecobg_db CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;')
    add_par(doc, "CREATE USER 'eco_user'@'localhost' IDENTIFIED BY 'eco_pass';")
    add_par(doc, "GRANT ALL PRIVILEGES ON ecobg_db.* TO 'eco_user'@'localhost';")
    add_par(doc, 'FLUSH PRIVILEGES;')
    add_par(doc, 'Шаг 2. Заполните файл .env:', True)
    add_par(doc, 'DATABASE_URL=mysql+pymysql://eco_user:eco_pass@localhost:3306/ecobg_db?charset=utf8mb4')
    add_par(doc, 'Шаг 3. Установите зависимости:', True)
    add_par(doc, 'python -m pip install -r requirements.txt')
    add_par(doc, 'Шаг 4. Инициализируйте БД (первый раз в чистой базе):', True)
    add_par(doc, '$env:PYTHONPATH=(Get-Location).Path')
    add_par(doc, 'python .\\scripts\\init_db.py')
    add_par(doc, 'python .\\scripts\\seed_basics.py')
    add_par(doc, 'python .\\scripts\\seed_help.py')
    add_par(doc, 'Шаг 5. Запуск сервера разработки:', True)
    add_par(doc, 'python .\\main.py')

    add_heading(doc, '5. Роли и доступы', 1)
    add_list(doc, [
        'Пользователь — создаёт объявления, инициирует сделки',
        'Менеджер — модерация объявлений и сделок, отчёты',
        'Администратор — пользователи, справочники, справка'
    ])

    add_heading(doc, '6. Основные сценарии использования', 1)
    add_list(doc, [
        'Регистрация и вход',
        'Создание объявления (передача/обмен/продажа)',
        'Просмотр карточки, переписка (в планах), инициирование обмена',
        'Подтверждение пожертвования владельцем',
        'Отметка переработки (метод, локация, комментарий)',
        'Личный кабинет: список своих объявлений и операций',
        'Справка: единый текст с подсказками, редактируемый администратором',
    ])

    add_heading(doc, '7. REST API (минимум)', 1)
    add_list(doc, [
        'GET /api/categories — список категорий',
        'GET /api/items — список объявлений (фильтры: category_id, status)',
        'POST /api/items — создание объявления (JSON)'
    ])

    add_heading(doc, '8. Интерфейс и навигация', 1)
    add_list(doc, [
        'Главная страница — hero-блок и последние объявления',
        'Каталог — фильтры (категория/статус), карточки',
        'Карточка — характеристики, действия (обмен/пожертвование/переработка)',
        'Личный кабинет — свои объявления и операции',
        'Справка — полезная информация по системе'
    ])

    add_heading(doc, '9. Типичные ошибки и их решение', 1)
    add_list(doc, [
        'Ошибка PyMySQL: требуется пакет cryptography — установите: python -m pip install cryptography',
        'Alembic не видит модели — при первой инициализации используйте scripts/init_db.py',
        'Нет доступа к БД — проверьте DATABASE_URL и права MySQL пользователя',
        'Зависимости — повторите установку: python -m pip install -r requirements.txt',
    ])

    add_heading(doc, '10. Замечания по разработке (ручной труд)', 1)
    add_par(doc, 'В проекте соблюдена структуризация слоёв и добавлены русские комментарии. Шаблоны и стили оформлены вручную с учётом UX. Документация (этот файл, user_guide.md, тест-план/кейсы) написана вручную. Рекомендуется вести частые понятные коммиты для демонстрации постепенной разработки.')

    add_heading(doc, '11. Планы развития', 1)
    add_list(doc, [
        'Расширение REST API (users, deals, messages)',
        'Документооборот: .docx акт, .xlsx отчёты',
        'PySide6 клиент (≥10 окон)',
        'Админ-панель (пользователи, отчёты, логи)'
    ])

    doc.add_page_break()
    add_heading(doc, 'Приложение A. Полезные команды (PowerShell)', 1)
    add_list(doc, [
        '$env:FLASK_APP = "main.py"',
        '$env:FLASK_ENV = "development"',
        'python -m pip install -r requirements.txt',
        'python .\\scripts\\init_db.py',
        'python .\\scripts\\seed_basics.py',
        'python .\\scripts\\seed_help.py',
        'python .\\main.py',
    ])

    out_path = 'docs/Guide_EcoBaraholka.docx'
    doc.save(out_path)
    print(f"[OK] Сформирован файл: {out_path}")


if __name__ == '__main__':
    main()
